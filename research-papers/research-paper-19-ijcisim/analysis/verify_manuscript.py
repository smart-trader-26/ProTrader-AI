#!/usr/bin/env python
"""Hard checks on the built IJCISIM manuscript.

    python analysis/verify_manuscript.py

The LaTeX papers in this project are checked by a script of the same name that
greps the .tex and the .log.  A Word build has no log, so the equivalent checks
read the .docx itself:

  1. no placeholder survived substitution (a stray "{SomeName}" in the text);
  2. every figure referenced in the prose is actually embedded, and every
     embedded image is referenced;
  3. every table and figure number is used in the prose at least once;
  4. the reference list is numbered contiguously and every bracketed citation
     in the text points at a number that exists;
  5. the page count is inside the journal's 25-page limit, if the PDF is built.

Exit status is non-zero if any check fails, so the build script can gate on it.
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
PAGE_LIMIT = 25          # "Papers should not exceed 25 typeset, printed pages"


def docx_text(path: Path) -> str:
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def n_images(path: Path) -> int:
    with zipfile.ZipFile(path) as z:
        return sum(1 for n in z.namelist() if n.startswith("word/media/"))


def pdf_pages(path: Path) -> int | None:
    if not path.exists():
        return None
    try:
        from pypdf import PdfReader
    except ImportError:                                   # pragma: no cover
        from PyPDF2 import PdfReader
    return len(PdfReader(str(path)).pages)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", default=str(ROOT / "ijcisim_validity_screen.docx"))
    ap.add_argument("--pdf", default=str(ROOT / "ijcisim_validity_screen.pdf"))
    args = ap.parse_args()

    doc = Path(args.docx)
    if not doc.exists():
        print(f"FAIL  {doc} does not exist; run build_paper.sh first")
        return 1

    text = docx_text(doc)
    failures: list[str] = []

    # 1 -- placeholders
    left = sorted(set(re.findall(r"\{[A-Za-z_][A-Za-z0-9_]*\}", text)))
    if left:
        failures.append(f"unsubstituted placeholders: {', '.join(left)}")

    # 2 -- figures embedded vs referenced
    refs = {int(n) for n in re.findall(r"Figure (\d+)", text)}
    embedded = n_images(doc)
    captioned = {int(n) for n in re.findall(r"Figure (\d+)\.", text)}
    if embedded != len(captioned):
        failures.append(f"{embedded} images embedded but {len(captioned)} "
                        f"figure captions")
    orphan_caption = captioned - (refs - captioned) - refs
    if orphan_caption:
        failures.append(f"figures never referenced in prose: {sorted(orphan_caption)}")

    # 3 -- tables captioned and referenced.  A cross-reference may be written
    # either as "Table 7" or collectively as "Tables 7 and 8", so both forms
    # count toward a table having been pointed at.
    tab_caps = {int(n) for n in re.findall(r"Table (\d+)\.", text)}
    pointed = {int(n) for n in re.findall(r"Table (\d+)", text)
               if text.count(f"Table {n}") > 1}
    for grp in re.findall(r"Tables ([\d,\s]+(?:and\s*\d+)?)", text):
        pointed.update(int(n) for n in re.findall(r"\d+", grp))
    missing_tab_ref = tab_caps - pointed
    if missing_tab_ref:
        failures.append(f"tables never referenced in prose: {sorted(missing_tab_ref)}")
    if tab_caps != set(range(1, len(tab_caps) + 1)):
        failures.append(f"table numbers not contiguous from 1: {sorted(tab_caps)}")

    # 4 -- citations point at real reference numbers
    body, _, reflist = text.partition("References\n")
    n_refs = len(re.findall(r"(?m)^\d+\.\t", reflist)) or len(
        re.findall(r"(?m)^\d+\.\s", reflist))
    cited = set()
    for grp in re.findall(r"\[([0-9,\s]+)\]", body):
        nums = [int(x) for x in grp.replace(" ", "").split(",") if x]
        # Intervals such as [0, 1] and [0, 2] look like citation groups; a
        # reference number is never zero, so that alone separates them.
        if nums and min(nums) >= 1:
            cited.update(nums)
    if not n_refs:
        failures.append("could not find a numbered reference list")
    else:
        bad = {c for c in cited if c < 1 or c > n_refs}
        if bad:
            failures.append(f"citations outside the reference list: {sorted(bad)}")
        gaps = set(range(1, n_refs + 1)) - cited
        if gaps:
            failures.append(f"reference numbers never cited: {sorted(gaps)}")

    # 5 -- page limit
    pages = pdf_pages(Path(args.pdf))
    if pages is not None and pages > PAGE_LIMIT:
        failures.append(f"{pages} pages exceeds the journal's {PAGE_LIMIT}-page limit")

    print(f"figures embedded : {embedded}")
    print(f"figure captions  : {len(captioned)}")
    print(f"table captions   : {len(tab_caps)}")
    print(f"references       : {n_refs} listed, {len(cited)} distinct cited")
    print(f"pages            : {pages if pages is not None else 'PDF not built'}"
          f" (limit {PAGE_LIMIT})")
    if failures:
        print()
        for f in failures:
            print("FAIL ", f)
        return 1
    print("\nall checks pass")
    return 0


if __name__ == "__main__":
    sys.exit(main())

"""Element builders for the IJCISIM manuscript, in python-docx.

IJCISIM (Cerebration Science Publishing, ISSN 2150-7988) does not distribute a
LaTeX class, and its submission checklist asks for a Microsoft Word file that is
single-spaced in a 10-point font with figures and tables placed in the text.
The manuscript is therefore built as a .docx rather than typeset.

There is no downloadable template either, so the formatting constants below were
read off a published article in the current house style (Vol.~16, 2024, the
Cerebration relaunch): A4, single column, Times New Roman throughout, a
first-page masthead, numbered section headings, run-in bold labels on the
abstract and keywords, figure captions beneath the figure and table captions
above the table, both introduced by a bold-italic "Figure N." / "Table N.", and
a numbered reference list.

Everything is direct formatting rather than named styles, which is what the
published articles themselves use and what survives a round trip through the
editor's own template.
"""
from __future__ import annotations

import os

import docx
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Cm, Pt, RGBColor

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

FONT = "Times New Roman"
BODY_PT = 10          # the checklist's "10-point font"
SMALL_PT = 9
TITLE_PT = 18
AUTHOR_PT = 12
HEAD_PT = 12
SUBHEAD_PT = 11
TABLE_PT = 9
INDENT_CM = 0.75


# --------------------------------------------------------------- primitives
def _xml(tag: str) -> str:
    return f'<w:{tag} xmlns:w="{W}"/>'


def _set_spacing(par, before=0, after=0, line=240, rule="auto"):
    pPr = par._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:spacing xmlns:w="{W}" w:before="{int(before * 20)}" '
        f'w:after="{int(after * 20)}" w:line="{line}" w:lineRule="{rule}"/>'))


def _keep_with_next(par):
    """Stop a caption from being orphaned from the float it introduces."""
    pPr = par._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepNext xmlns:w="{W}"/>'))


def _bottom_border(par, sz=6):
    pPr = par._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr xmlns:w="{W}"><w:bottom w:val="single" w:sz="{sz}" '
        f'w:space="1" w:color="000000"/></w:pBdr>'))


def _cell_border(cell, edge: str, sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders xmlns:w="{W}"><w:{edge} w:val="single" w:sz="{sz}" '
        f'w:space="0" w:color="000000"/></w:tcBorders>'))


def _run(par, text, *, size=BODY_PT, bold=False, italic=False,
         superscript=False, color=None):
    r = par.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if superscript:
        r.font.superscript = True
    if color:
        r.font.color.rgb = RGBColor(*color)
    rPr = r._r.get_or_add_rPr()
    rPr.append(parse_xml(
        f'<w:rFonts xmlns:w="{W}" w:ascii="{FONT}" w:hAnsi="{FONT}" '
        f'w:cs="{FONT}" w:eastAsia="{FONT}"/>'))
    return r


def rich(par, text, *, size=BODY_PT, bold=False, italic=False):
    """Append text with a tiny inline markup: **bold**, *italic*, ^superscript^.

    Keeping the markup this small means the manuscript source stays readable as
    plain text while still producing real Word runs rather than an image.
    """
    i = 0
    buf = ""
    while i < len(text):
        two = text[i:i + 2]
        one = text[i]
        if two == "**":
            if buf:
                _run(par, buf, size=size, bold=bold, italic=italic)
                buf = ""
            j = text.find("**", i + 2)
            j = len(text) if j < 0 else j
            _run(par, text[i + 2:j], size=size, bold=True, italic=italic)
            i = j + 2
        elif one in "*^":
            # An unmatched marker is a literal character, not an opener that
            # runs to the end of the paragraph and italicises everything after
            # it.  Use U+2217 for a mathematical star so tau-star cannot be
            # mistaken for markup in the first place.
            j = text.find(one, i + 1)
            if j < 0:
                buf += one
                i += 1
                continue
            if buf:
                _run(par, buf, size=size, bold=bold, italic=italic)
                buf = ""
            _run(par, text[i + 1:j], size=size, bold=bold,
                 italic=italic or one == "*", superscript=(one == "^"))
            i = j + 1
        else:
            buf += one
            i += 1
    if buf:
        _run(par, buf, size=size, bold=bold, italic=italic)


# ------------------------------------------------------------------ document
def new_document():
    doc = docx.Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
    sec.left_margin = sec.right_margin = Cm(3.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.different_first_page_header_footer = True

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BODY_PT)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.space_before = Pt(0)
    rPr = st.element.get_or_add_rPr()
    rPr.append(parse_xml(
        f'<w:rFonts xmlns:w="{W}" w:ascii="{FONT}" w:hAnsi="{FONT}" '
        f'w:cs="{FONT}" w:eastAsia="{FONT}"/>'))
    return doc


def page_numbers(doc):
    """Centred page number in the footer, as the published articles carry."""
    fields = [f'<w:fldChar xmlns:w="{W}" w:fldCharType="begin"/>',
              f'<w:instrText xmlns:w="{W}" xml:space="preserve"> PAGE </w:instrText>',
              f'<w:fldChar xmlns:w="{W}" w:fldCharType="end"/>']
    for footer in (doc.sections[0].footer, doc.sections[0].first_page_footer):
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.font.name = FONT
        r.font.size = Pt(SMALL_PT)
        for el in fields:
            r._r.append(parse_xml(el))


# ---------------------------------------------------------------- front page
def masthead(doc, *, journal, issn, volume, year, publisher):
    p = doc.add_paragraph()
    _set_spacing(p, after=0)
    _run(p, journal, size=SMALL_PT)
    p2 = doc.add_paragraph()
    _set_spacing(p2, after=0)
    _run(p2, f"ISSN {issn} Volume {volume} ({year})", size=SMALL_PT)
    p3 = doc.add_paragraph()
    _set_spacing(p3, after=6)
    _bottom_border(p3)
    _run(p3, f"© {publisher}", size=SMALL_PT)

    lab = doc.add_paragraph()
    _set_spacing(lab, before=6, after=12)
    _run(lab, "Article", size=SMALL_PT)


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=6, after=12, line=280)
    _run(p, text, size=TITLE_PT, bold=True)


def authors(doc, entries):
    """entries: list of (name, marker) where marker is e.g. '1,*' or '2'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=6, after=8)
    for k, (name, marker) in enumerate(entries):
        if k:
            _run(p, " and ", size=AUTHOR_PT, bold=True)
        _run(p, name, size=AUTHOR_PT, bold=True)
        _run(p, " " + marker, size=AUTHOR_PT, bold=True, superscript=True)


def affiliations(doc, lines, correspondence):
    for k, line in enumerate(lines, start=1):
        p = doc.add_paragraph()
        _set_spacing(p, after=0)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        _run(p, f"{k}\t", size=SMALL_PT)
        _run(p, line, size=SMALL_PT)
    p = doc.add_paragraph()
    _set_spacing(p, after=10)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    _run(p, "*\t", size=SMALL_PT)
    _run(p, correspondence, size=SMALL_PT)


def abstract(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, before=8, after=8, line=280)
    _run(p, "Abstract: ", size=BODY_PT, bold=True)
    rich(p, text, size=BODY_PT)


def keywords(doc, words):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, before=4, after=10)
    _bottom_border(p)
    _run(p, "Keywords: ", size=BODY_PT, bold=True)
    _run(p, "; ".join(words), size=BODY_PT)


# ------------------------------------------------------------------ sections
def heading(doc, text):
    p = doc.add_paragraph()
    _set_spacing(p, before=12, after=6)
    _run(p, text, size=HEAD_PT, bold=True)


def subheading(doc, text):
    p = doc.add_paragraph()
    _set_spacing(p, before=8, after=4)
    _run(p, text, size=SUBHEAD_PT, bold=True, italic=True)


def body(doc, text, *, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, after=0, line=240)
    if indent:
        p.paragraph_format.first_line_indent = Cm(INDENT_CM)
    rich(p, text)
    return p


def listing(doc, items, *, numbered=True, start=1):
    for k, it in enumerate(items, start=start):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(p, before=2, after=2)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        _run(p, f"({k})\t" if numbered else "•\t")
        rich(p, it)


def equation(doc, omml_xml, number):
    """Centred display equation with a right-aligned number."""
    from omml import NS
    p = doc.add_paragraph()
    _set_spacing(p, before=6, after=6)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:tabs xmlns:w="{W}"><w:tab w:val="center" w:pos="4200"/>'
        f'<w:tab w:val="right" w:pos="8400"/></w:tabs>'))
    p.add_run("\t")
    p._p.append(parse_xml(f"<m:oMath {NS}>{omml_xml}</m:oMath>"))
    _run(p, "\t" + number)
    return p


def figure(doc, path, number, caption, width_in=4.6):
    from docx.shared import Inches
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=10, after=4)
    _keep_with_next(p)
    p.add_run().add_picture(path, width=Inches(width_in))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(c, after=10)
    _run(c, f"Figure {number}. ", size=SMALL_PT, bold=True, italic=True)
    rich(c, caption, size=SMALL_PT)


def table(doc, spec, number):
    """Render one table from the JSON spec written by 06b_tables_json.py."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(cap, before=10, after=4)
    _keep_with_next(cap)
    _run(cap, f"Table {number}. ", size=SMALL_PT, bold=True, italic=True)
    rich(cap, spec["caption"], size=SMALL_PT)

    header = spec["header"]
    ncol = len(header)
    rows = spec["rows"]
    t = doc.add_table(rows=1 + len(rows), cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Word's autofit narrows the label column until multi-word row labels wrap
    # onto three lines, so the layout is fixed and the widths are declared.
    # Default: a wide first column for the labels, the rest shared evenly.
    t.autofit = False
    t._tbl.tblPr.append(parse_xml(
        f'<w:tblLayout xmlns:w="{W}" w:type="fixed"/>'))
    sec = doc.sections[0]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    fracs = spec.get("widths")
    if not fracs:
        first = 0.30 if ncol > 3 else 0.28
        fracs = [first] + [(1 - first) / (ncol - 1)] * (ncol - 1)
    widths = [int(usable * f) for f in fracs]
    for row in t.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]

    show_header = any(h.strip() for h in header)
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        _set_spacing(cell.paragraphs[0], before=2, after=2)
        if show_header:
            _run(cell.paragraphs[0], h, size=TABLE_PT, bold=False)
        _cell_border(cell, "top")
        _cell_border(cell, "bottom")

    align = spec.get("align", ["l"] + ["c"] * (ncol - 1))
    for i, row in enumerate(rows):
        cells = t.rows[1 + i].cells
        if len(row) == 1:                       # a panel label spanning the row
            merged = cells[0]
            for c in cells[1:]:
                merged = merged.merge(c)
            par = merged.paragraphs[0]
            _set_spacing(par, before=4, after=2)
            _run(par, row[0], size=TABLE_PT, italic=True)
            continue
        bold = row[0] in ("All", "Pooled", "Total fitted objects",
                          "Always-up base rate")
        for j in range(ncol):
            txt = row[j] if j < len(row) else ""
            par = cells[j].paragraphs[0]
            par.alignment = {
                "l": WD_ALIGN_PARAGRAPH.LEFT,
                "c": WD_ALIGN_PARAGRAPH.CENTER,
                "r": WD_ALIGN_PARAGRAPH.RIGHT}[align[j] if j < len(align) else "c"]
            _set_spacing(par, before=1, after=1)
            _run(par, txt, size=TABLE_PT, bold=bold)
        if i == len(rows) - 1:
            for c in cells:
                _cell_border(c, "bottom")

    if spec.get("note"):
        n = doc.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(n, before=3, after=10)
        rich(n, spec["note"], size=SMALL_PT)
    else:
        _set_spacing(doc.add_paragraph(), after=8)


def references(doc, entries):
    for k, e in enumerate(entries, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(p, before=2, after=2)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        _run(p, f"{k}.\t", size=SMALL_PT)
        rich(p, e, size=SMALL_PT)


def backmatter_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, before=2, after=6)
    rich(p, text, size=SMALL_PT)


def copyright_footer(doc, text):
    p = doc.add_paragraph()
    _set_spacing(p, before=14, after=0)
    _bottom_border(p, sz=4)
    _run(p, "", size=SMALL_PT)
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(q, before=4, after=0)
    _run(q, text, size=SMALL_PT)


# ---------------------------------------------------------------- conversion
def to_pdf(docx_path: str, pdf_path: str):
    """Convert through Word itself, so the PDF is what an editor would see."""
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    d = word.Documents.Open(os.path.abspath(docx_path))
    d.Fields.Update()
    d.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)      # wdFormatPDF
    pages = d.ComputeStatistics(2)                            # wdStatisticPages
    words = d.ComputeStatistics(0)                            # wdStatisticWords
    d.Close(False)
    word.Quit()
    print(f"wrote {pdf_path}: {pages} pages, {words:,} words")
    return pages

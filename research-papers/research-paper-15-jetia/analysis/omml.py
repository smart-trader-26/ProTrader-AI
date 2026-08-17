"""Minimal Office MathML (OMML) builders.

Word stores equations as OMML, which is what the Equation Editor produces.
Building OMML directly means the equations in the generated .doc are real,
editable Word equations rather than pictures.
"""
from __future__ import annotations

from docx.oxml.ns import nsmap, qn
from docx.oxml.parser import parse_xml

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
nsmap.setdefault("m", M)
NS = f'xmlns:m="{M}" xmlns:w="{nsmap["w"]}"'


def _esc(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def r(text: str, *, sty: str = "i") -> str:
    """A math run.  sty: 'i' italic (variables), 'p' upright (operators)."""
    pr = "" if sty == "i" else '<m:rPr><m:sty m:val="p"/></m:rPr>'
    return f"<m:r>{pr}<m:t xml:space='preserve'>{_esc(text)}</m:t></m:r>"


def up(text: str) -> str:
    return r(text, sty="p")


def sub(base: str, s: str) -> str:
    return (f"<m:sSub><m:e>{base}</m:e><m:sub>{s}</m:sub></m:sSub>")


def sup(base: str, s: str) -> str:
    return (f"<m:sSup><m:e>{base}</m:e><m:sup>{s}</m:sup></m:sSup>")


def subsup(base: str, s: str, p: str) -> str:
    return (f"<m:sSubSup><m:e>{base}</m:e><m:sub>{s}</m:sub>"
            f"<m:sup>{p}</m:sup></m:sSubSup>")


def frac(num: str, den: str) -> str:
    return (f"<m:f><m:fPr><m:type m:val='bar'/></m:fPr>"
            f"<m:num>{num}</m:num><m:den>{den}</m:den></m:f>")


def rad(deg: str, e: str) -> str:
    dp = "<m:degHide m:val='1'/>" if not deg else ""
    return (f"<m:rad><m:radPr>{dp}</m:radPr><m:deg>{deg}</m:deg>"
            f"<m:e>{e}</m:e></m:rad>")


def nary(op: str, lo: str, hi: str, e: str) -> str:
    return (f"<m:nary><m:naryPr><m:chr m:val='{op}'/>"
            f"<m:limLoc m:val='undOvr'/></m:naryPr>"
            f"<m:sub>{lo}</m:sub><m:sup>{hi}</m:sup><m:e>{e}</m:e></m:nary>")


def hat(e: str) -> str:
    return (f"<m:acc><m:accPr><m:chr m:val='&#770;'/></m:accPr>"
            f"<m:e>{e}</m:e></m:acc>")


def tilde(e: str) -> str:
    return (f"<m:acc><m:accPr><m:chr m:val='&#771;'/></m:accPr>"
            f"<m:e>{e}</m:e></m:acc>")


def d(*parts: str) -> str:
    """Concatenate OMML fragments."""
    return "".join(parts)


def paren(inner: str, left: str = "(", right: str = ")") -> str:
    return (f"<m:d><m:dPr><m:begChr m:val='{left}'/>"
            f"<m:endChr m:val='{right}'/></m:dPr><m:e>{inner}</m:e></m:d>")


def inline(par, omml: str) -> None:
    """Append an inline equation to an existing python-docx paragraph."""
    par._p.append(parse_xml(f"<m:oMath {NS}>{omml}</m:oMath>"))


def display(doc, omml: str, number: str, style: str = "Text",
            tab_pos_twips: int = 5000):
    """Centred display equation with a right-aligned number in brackets."""
    p = doc.add_paragraph(style=style)
    pPr = p._p.get_or_add_pPr()
    tabs = parse_xml(
        f'<w:tabs xmlns:w="{nsmap["w"]}">'
        f'<w:tab w:val="center" w:pos="{tab_pos_twips // 2}"/>'
        f'<w:tab w:val="right" w:pos="{tab_pos_twips}"/></w:tabs>')
    pPr.append(tabs)
    p.add_run("\t")
    p._p.append(parse_xml(f"<m:oMath {NS}>{omml}</m:oMath>"))
    p.add_run("\t" + number)
    return p


__all__ = ["r", "up", "sub", "sup", "subsup", "frac", "rad", "nary", "hat",
           "tilde", "d", "paren", "inline", "display", "qn"]

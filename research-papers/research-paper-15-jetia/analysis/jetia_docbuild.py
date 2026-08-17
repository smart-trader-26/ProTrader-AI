"""Shared helpers for writing the ITEGAM-JETIA manuscript with python-docx.

The manuscript is produced by editing a copy of the journal's own template
(Template.docx, downloaded from itegam-jetia.org).  That template carries no
named paragraph styles: every element is direct formatting.  The paragraph
and run property blocks below were therefore lifted verbatim out of the
template's own paragraphs, so an element built here is byte-for-byte
formatted the way the corresponding element of the template is.

Element                     lifted from template paragraph
--------------------------  ---------------------------------
section heading             "I. INTRODUCTION"
first-level subheading      "II.1 SUBTITLE"
second-level subheading     "II.1.1 Materials and Methods"
body text                   "It must define the problem ..."
figure / caption / source   "Figure 1: Figure title." block
table title / source        "Table 1: Table title." block
equation                    the quadratic-formula paragraph
reference entry             "[1] J. Nogueira, ..."
"""
from __future__ import annotations

import os
import re

import docx
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Inches, Pt

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ------------------------------------------------------------------ formats
TNR = ('<w:rFonts w:ascii="Times New Roman" w:eastAsia="Times New Roman" '
       'w:hAnsi="Times New Roman" w:cs="Times New Roman"/>')
TNR2 = '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
BLUE = '<w:color w:val="1F4E79" w:themeColor="accent1" w:themeShade="80"/>'
SZ10 = '<w:sz w:val="20"/><w:szCs w:val="20"/>'
SZ8 = '<w:sz w:val="16"/><w:szCs w:val="16"/>'
EN = '<w:lang w:val="en-US" w:eastAsia="pt-BR"/>'

# section heading: centred, bold, dark blue, 10 pt, Normal's 8 pt space after
HEAD_P = '<w:spacing w:line="240" w:lineRule="auto"/><w:jc w:val="center"/>'
HEAD_R = TNR2 + '<w:b/><w:bCs/>' + BLUE + SZ10 + EN

# first-level subheading (II.1 ...): left, bold, dark blue, upper case
SUB1_P = '<w:spacing w:line="240" w:lineRule="auto"/>'
SUB1_R = TNR2 + '<w:b/><w:bCs/><w:iCs/>' + BLUE + SZ10 + '<w:lang w:val="en-US"/>'

# second-level subheading (II.1.1 ...): left, regular, dark blue
SUB2_P = '<w:spacing w:line="240" w:lineRule="auto"/><w:jc w:val="both"/>'
SUB2_R = TNR2 + '<w:iCs/>' + BLUE + SZ10 + '<w:lang w:val="en-US"/>'

# body paragraph: justified, 1 cm first-line indent, no space after
BODY_P = ('<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
          '<w:ind w:firstLine="567"/><w:jc w:val="both"/>')
BODY_R = TNR + '<w:color w:val="222222"/>' + SZ10 + EN

# figure image, caption and source: centred, no space after
FIGIMG_P = ('<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
            '<w:ind w:right="212"/><w:jc w:val="center"/>')
CAP_P = ('<w:shd w:val="clear" w:color="auto" w:fill="FFFFFF"/>'
         '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
         '<w:ind w:right="212"/><w:jc w:val="center"/>')
CAP_R = TNR + SZ10 + EN

# table title (above the table) and source line (below it)
TABTITLE_P = ('<w:shd w:val="clear" w:color="auto" w:fill="FFFFFF"/>'
              '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
              '<w:jc w:val="center"/>')
TABSRC_P = ('<w:shd w:val="clear" w:color="auto" w:fill="FFFFFF"/>'
            '<w:spacing w:before="20" w:after="0" w:line="240" '
            'w:lineRule="auto"/><w:ind w:right="210"/>'
            '<w:jc w:val="center"/>')

# display equation: the number sits on a right tab at the text margin
EQ_P = ('<w:tabs><w:tab w:val="center" w:pos="5100"/>'
        '<w:tab w:val="right" w:pos="10800"/></w:tabs>'
        '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
        '<w:ind w:firstLine="0"/><w:jc w:val="both"/>')
EQ_R = TNR2 + '<w:bCs/><w:iCs/><w:color w:val="000000"/>' + SZ10

# reference entry: justified, 8 pt
REF_P = ('<w:tabs><w:tab w:val="left" w:pos="567"/></w:tabs>'
         '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
         '<w:jc w:val="both"/>')
REF_R = TNR2 + '<w:bCs/><w:iCs/><w:color w:val="000000"/>' + SZ8 + \
    '<w:lang w:val="en-US"/>'

# blank spacer paragraph, as the template puts one before each heading
GAP_P = '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'

SOURCE_LINE = "Source: Authors, (2026)."

# lightweight inline markup so body prose can carry real subscripts,
# superscripts and italic variables instead of ASCII stand-ins
RICH = re.compile(r"\{([_^ib]):?([^}]*)\}")


def load_template(path: str) -> docx.Document:
    return docx.Document(path)


def _ppr(p, xml: str) -> None:
    """Replace a paragraph's properties with the given pPr body."""
    old = p._p.find(qn("w:pPr"))
    if old is not None:
        p._p.remove(old)
    p._p.insert(0, parse_xml(f'<w:pPr xmlns:w="{W}">{xml}</w:pPr>'))


def _rpr(run, xml: str) -> None:
    old = run._r.find(qn("w:rPr"))
    if old is not None:
        run._r.remove(old)
    run._r.insert(0, parse_xml(f'<w:rPr xmlns:w="{W}">{xml}</w:rPr>'))


def rich_run(par, text: str, rpr: str) -> None:
    """Add runs to par, honouring {_sub}, {^sup} and {i:italic} markup."""
    def emit(chunk, kind=None):
        if not chunk:
            return
        r = par.add_run(chunk)
        extra = ""
        if kind == "_":
            extra = "<w:vertAlign w:val=\"subscript\"/>"
        elif kind == "^":
            extra = "<w:vertAlign w:val=\"superscript\"/>"
        elif kind == "i" and "<w:i/>" not in rpr:
            # only add italic when the base format is not italic already:
            # a duplicate <w:i/> inside one rPr is not schema-valid
            extra = "<w:i/><w:iCs/>"
        elif kind == "b" and "<w:b/>" not in rpr:
            extra = "<w:b/><w:bCs/>"
        _rpr(r, rpr + extra)

    pos = 0
    for m in RICH.finditer(text):
        emit(text[pos:m.start()])
        emit(m.group(2), m.group(1))
        pos = m.end()
    emit(text[pos:])


def strip_rich(text: str) -> str:
    return RICH.sub(lambda m: m.group(2), text)


def clear_after_title_block(doc) -> int:
    """Drop the template's specimen body, keeping the title block.

    The body is: title table, an empty paragraph, a paragraph carrying the
    first section break, then the specimen article, then the body-level
    section properties.  Everything between the section break and those
    properties is removed so the real manuscript can be appended.
    """
    body = doc.element.body
    kids = list(body)
    cut = None
    for i, el in enumerate(kids):
        if el.tag == qn("w:p"):
            pPr = el.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                cut = i
                break
    if cut is None:
        raise RuntimeError("section break paragraph not found in template")
    removed = 0
    for el in kids[cut + 1:]:
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)
        removed += 1
    return removed


class Builder:
    """Appends manuscript content before the body-level section properties."""

    def __init__(self, doc):
        self.doc = doc
        self.body = doc.element.body
        self.anchor = self.body.find(qn("w:sectPr"))
        if self.anchor is None:
            raise RuntimeError("body sectPr not found")
        self.fig_no = 0
        self.tab_no = 0
        self.eq_no = 0

    # ---------------------------------------------------------- primitives
    def _add(self, ppr_xml: str):
        p = self.doc.add_paragraph()
        self.body.remove(p._p)
        self.anchor.addprevious(p._p)
        _ppr(p, ppr_xml)
        return p

    def gap(self):
        return self._add(GAP_P)

    # ------------------------------------------------------------ headings
    def head(self, text: str):
        """Centred section heading, e.g. 'I. INTRODUCTION'."""
        self.gap()
        p = self._add(HEAD_P)
        r = p.add_run(text.upper())
        _rpr(r, HEAD_R)
        return p

    def sub1(self, text: str):
        """First-level subheading, e.g. 'II.1 NEWS SENTIMENT AND PRICES'."""
        self.gap()
        p = self._add(SUB1_P)
        r = p.add_run(text.upper())
        _rpr(r, SUB1_R)
        return p

    def sub2(self, text: str):
        """Second-level subheading, set in sentence case like the template."""
        self.gap()
        p = self._add(SUB2_P)
        rich_run(p, text, SUB2_R)
        return p

    # ---------------------------------------------------------------- text
    def text(self, text: str, *, italic: bool = False):
        p = self._add(BODY_P)
        rich_run(p, text, BODY_R + ("<w:i/><w:iCs/>" if italic else ""))
        return p

    # ----------------------------------------------------------- equations
    def equation(self, omml: str):
        """Centred display equation with a right-aligned number."""
        from omml import NS

        self.eq_no += 1
        p = self._add(EQ_P)
        r = p.add_run("\t")
        _rpr(r, EQ_R)
        p._p.append(parse_xml(f"<m:oMath {NS}>{omml}</m:oMath>"))
        r = p.add_run(f"\t({self.eq_no})")
        _rpr(r, EQ_R)
        return self.eq_no

    # ------------------------------------------------------------- figures
    def figure(self, path: str, caption: str, width_in: float = 4.6):
        """Image, then 'Figure n: caption' and the source line beneath it."""
        self.gap()
        self.fig_no += 1
        p = self._add(FIGIMG_P)
        r = p.add_run()
        _rpr(r, TNR2 + '<w:noProof/>' + SZ10)
        r.add_picture(path, width=Inches(width_in))
        c = self._add(CAP_P)
        rc = c.add_run(f"Figure {self.fig_no}: {caption}")
        _rpr(rc, CAP_R)
        s = self._add(CAP_P)
        rs = s.add_run(SOURCE_LINE)
        _rpr(rs, CAP_R)
        return self.fig_no

    # -------------------------------------------------------------- tables
    def table(self, title: str, rows: list[list[str]], font_pt: float = 9.0):
        """Caption above the table and the source line below, as the
        template requires."""
        self.gap()
        self.tab_no += 1
        cap = self._add(TABTITLE_P)
        rc = cap.add_run(f"Table {self.tab_no}: {title}")
        _rpr(rc, CAP_R)

        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        t.autofit = True
        self.body.remove(t._tbl)
        self.anchor.addprevious(t._tbl)
        # a long table may break across pages, so repeat the column names
        hdr = t.rows[0]._tr.get_or_add_trPr()
        hdr.append(parse_xml(f'<w:tblHeader xmlns:w="{W}"/>'))
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                c = t.cell(i, j)
                par = c.paragraphs[0]
                _ppr(par, '<w:spacing w:after="0" w:line="240" '
                          'w:lineRule="auto"/>'
                          + ('<w:jc w:val="left"/>' if j == 0
                             else '<w:jc w:val="center"/>'))
                run = par.add_run(str(cell))
                _rpr(run, TNR + ('<w:b/>' if i == 0 else "")
                     + f'<w:sz w:val="{int(font_pt * 2)}"/>'
                       f'<w:szCs w:val="{int(font_pt * 2)}"/>' + EN)
        src = self._add(TABSRC_P)
        rs = src.add_run(SOURCE_LINE)
        _rpr(rs, CAP_R)
        return self.tab_no

    # ---------------------------------------------------------- references
    def ref(self, n: int, text: str):
        p = self._add(REF_P)
        r = p.add_run(f"[{n}] {text}")
        _rpr(r, REF_R)
        self._add(REF_P)          # blank line between entries, as in template
        return p


# --------------------------------------------------------------- title block
def _copy_fmt(run, rpr_xml, vert: str = ""):
    if not rpr_xml:
        return
    xml = rpr_xml
    if vert:
        xml = xml.replace(
            "</w:rPr>", f'<w:vertAlign w:val="{vert}"/></w:rPr>')
    old = run._r.find(qn("w:rPr"))
    if old is not None:
        run._r.remove(old)
    run._r.insert(0, parse_xml(xml))


def _emit_cell_runs(par, text: str, rpr_xml) -> None:
    """Write text into a title-block paragraph, honouring {^sup}/{_sub}."""
    pos = 0
    for m in RICH.finditer(text):
        if m.start() > pos:
            _copy_fmt(par.add_run(text[pos:m.start()]), rpr_xml)
        vert = {"^": "superscript", "_": "subscript"}.get(m.group(1), "")
        _copy_fmt(par.add_run(m.group(2)), rpr_xml, vert)
        pos = m.end()
    if pos < len(text):
        _copy_fmt(par.add_run(text[pos:]), rpr_xml)


def set_cell(cell, text) -> None:
    """Replace a title-block cell's text, keeping its first run's format.

    `text` may be a string or a list of strings, one per line; extra lines
    reuse the paragraph format of the line above them.
    """
    from lxml import etree

    lines = [text] if isinstance(text, str) else list(text)
    ps = cell.paragraphs
    first = ps[0]
    rpr_xml = None
    if first.runs:
        el = first.runs[0]._r.find(qn("w:rPr"))
        if el is not None:
            rpr_xml = etree.tostring(el, encoding="unicode")
    ppr_xml = None
    el = first._p.find(qn("w:pPr"))
    if el is not None:
        ppr_xml = etree.tostring(el, encoding="unicode")

    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    for extra in ps[1:]:
        extra._p.getparent().remove(extra._p)

    _emit_cell_runs(first, lines[0], rpr_xml)
    prev = first._p
    for line in lines[1:]:
        new_p = parse_xml(f'<w:p xmlns:w="{W}">{ppr_xml or ""}</w:p>')
        prev.addnext(new_p)
        prev = new_p
        par = docx.text.paragraph.Paragraph(new_p, cell)
        _emit_cell_runs(par, line, rpr_xml)


def drop_rows(table, indices) -> None:
    """Remove title-block rows (used to blind the manuscript)."""
    for i in sorted(indices, reverse=True):
        tr = table.rows[i]._tr
        tr.getparent().remove(tr)


def set_keywords(cell, words: list[str]) -> None:
    """Fill the 'Keywords:' cell: one term per line, as the template shows."""
    ps = cell.paragraphs
    label, slots = ps[0], ps[1:]
    fmt = None
    if slots and slots[0].runs:
        from lxml import etree
        el = slots[0].runs[0]._r.find(qn("w:rPr"))
        if el is not None:
            fmt = etree.tostring(el, encoding="unicode")
    del label                                    # 'Keywords:' stays as it is
    for i, par in enumerate(slots):
        for r in list(par.runs):
            r._element.getparent().remove(r._element)
        if i < len(words):
            sep = "." if i == len(words) - 1 else ","
            run = par.add_run(words[i] + sep)
            if fmt:
                old = run._r.find(qn("w:rPr"))
                if old is not None:
                    run._r.remove(old)
                run._r.insert(0, parse_xml(fmt))
    for par in slots[len(words):]:
        par._p.getparent().remove(par._p)


PLACEHOLDER_NAMES = "One, Two and Three"


def set_running_header(doc, text: str) -> int:
    """Rewrite the author-name placeholder in the pages 2+ running header.

    That header reads '<names>, ITEGAM-JETIA, Manaus, v.XX ...' and lives
    inside a text box, so the runs have to be reached through the header
    part's element tree rather than through its paragraphs.  The
    volume/issue and DOI fields are yellow-highlighted, i.e. filled in by
    the journal, and are left untouched.  Passing an empty string removes
    the names and the comma that followed them, which is how the running
    header of a published JETIA article reads.
    """
    hits = 0
    parts = []
    for section in doc.sections:
        for attr in ("header", "first_page_header", "even_page_header"):
            part = getattr(section, attr, None)
            if part is not None and part not in parts:
                parts.append(part)
    for part in parts:
        runs = list(part._element.iter(qn("w:r")))
        for i, r in enumerate(runs):
            ts = r.findall(qn("w:t"))
            if len(ts) != 1 or (ts[0].text or "") != PLACEHOLDER_NAMES:
                continue
            if text:
                ts[0].text = text
            else:
                ts[0].text = ""
                if i + 1 < len(runs):
                    nxt = runs[i + 1].findall(qn("w:t"))
                    if len(nxt) == 1 and (nxt[0].text or "").strip() == ",":
                        nxt[0].text = ""
            hits += 1
    return hits


def scrub_properties(doc) -> None:
    """Remove author identification from the file properties, which the
    journal's submission checklist requires for blind review."""
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.comments = ""
    cp.category = ""
    cp.keywords = ""
    cp.identifier = ""


def to_doc(docx_path: str, doc_path: str):
    """Convert to Word 97-2003 .doc and report page and word counts."""
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    d = word.Documents.Open(os.path.abspath(docx_path))
    d.SaveAs2(os.path.abspath(doc_path), FileFormat=0)     # wdFormatDocument
    pages = d.ComputeStatistics(2)
    words = d.ComputeStatistics(0)
    d.Close(False)
    word.Quit()
    print(f"wrote {doc_path}: {pages} pages, {words} words")
    return pages, words


def to_pdf(docx_path: str, pdf_path: str):
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    d = word.Documents.Open(os.path.abspath(docx_path))
    d.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)    # wdFormatPDF
    pages = d.ComputeStatistics(2)
    d.Close(False)
    word.Quit()
    print(f"wrote {pdf_path}: {pages} pages")
    return pages

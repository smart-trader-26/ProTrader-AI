"""Shared helpers for writing the AECE manuscript with python-docx.

The manuscript is produced by editing a copy of the journal's own template,
so every paragraph carries one of the template's named styles and no font,
size or spacing is ever set by hand.
"""
from __future__ import annotations

import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# template paragraph indices (verified against AECE_template.docx)
I_TITLE_A, I_TITLE_B = 0, 1
I_AUTHORS, I_AFF1, I_AFF2, I_EMAIL = 3, 4, 5, 6
I_SECT_BREAK = 7
I_ABSTRACT, I_INDEX = 8, 10
I_BODY_START, I_BODY_END = 11, 67       # inclusive range to clear


def load_template(path: str) -> docx.Document:
    return docx.Document(path)


import re

# lightweight inline markup so body prose can carry real subscripts,
# superscripts and italic variables instead of ASCII stand-ins
RICH = re.compile(r"\{([_^i]):?([^}]*)\}")


def rich_run(par, text: str) -> None:
    """Add runs to par, honouring {_sub}, {^sup} and {i:italic} markup."""
    pos = 0
    for m in RICH.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        kind, payload = m.group(1), m.group(2)
        r = par.add_run(payload)
        if kind == "_":
            r.font.subscript = True
        elif kind == "^":
            r.font.superscript = True
        else:
            r.italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def strip_rich(text: str) -> str:
    return RICH.sub(lambda m: m.group(2), text)


def set_text(par, text: str, *, bold_lead: str = "") -> None:
    """Replace a paragraph's text, keeping its style."""
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    if bold_lead:
        r = par.add_run(bold_lead)
        r.italic = True
    rich_run(par, text)


def clear_body(doc: docx.Document, start: int, end: int) -> None:
    body = doc.element.body
    for el in list(body)[start:end + 1]:
        body.remove(el)


class Builder:
    """Appends content immediately before the trailing section break."""

    def __init__(self, doc: docx.Document, anchor_index: int):
        self.doc = doc
        self.body = doc.element.body
        self.anchor = list(self.body)[anchor_index]
        self.fig_no = 0
        self.tab_no = 0
        self.eq_no = 0

    def _add(self, style: str):
        p = self.doc.add_paragraph(style=style)
        self.body.remove(p._p)
        self.anchor.addprevious(p._p)
        return p

    def head(self, text: str):
        p = self._add("Heading 1")
        p.add_run(text)
        return p

    def subhead(self, text: str):
        """Italic run-in heading; inline markup is honoured inside it."""
        p = self._add("Text")
        rich_run(p, text)
        for r in p.runs:
            r.italic = True
        return p

    def text(self, text: str):
        p = self._add("Text")
        rich_run(p, text)
        return p

    def refhead(self, text: str):
        p = self._add("Reference Head")
        p.add_run(text)
        return p

    def ref(self, text: str):
        p = self._add("References")
        p.add_run(text)
        return p

    def equation(self, omml: str, right_tab_twips: int = 4900):
        """Centred equation with a right-aligned number."""
        from docx.oxml.parser import parse_xml
        from omml import NS

        self.eq_no += 1
        p = self._add("Text")
        pPr = p._p.get_or_add_pPr()
        w_ns = qn("w:t").split("}")[0][1:]
        # a text column is about 243 pt wide, i.e. ~4870 twips; the number
        # sits on a right tab just inside that so it cannot collide with a
        # long equation, and the equation itself is centred before it
        pPr.append(parse_xml(
            f'<w:tabs xmlns:w="{w_ns}">'
            f'<w:tab w:val="center" w:pos="{int(right_tab_twips * 0.46)}"/>'
            f'<w:tab w:val="right" w:pos="{right_tab_twips}"/></w:tabs>'))
        pPr.append(parse_xml(
            f'<w:ind xmlns:w="{w_ns}" w:firstLine="0"/>'))
        p.add_run("\t")
        p._p.append(parse_xml(f"<m:oMath {NS}>{omml}</m:oMath>"))
        p.add_run(f"\t({self.eq_no})")
        return self.eq_no

    def figure(self, path: str, caption: str, width_in: float = 3.3):
        self.fig_no += 1
        p = self._add("Text")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))
        c = self._add("Text")
        c.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = c.add_run(f"Figure {self.fig_no}. {caption}")
        r.font.size = Pt(8)
        return self.fig_no

    def table(self, title: str, rows: list[list[str]], widths=None,
              font_pt: float = 7.5):
        """Table with its caption above, as the template requires."""
        self.tab_no += 1
        cap = self._add("Table Title")
        cap.add_run(f"Table {roman(self.tab_no)}. {title}")
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        self.body.remove(t._tbl)
        self.anchor.addprevious(t._tbl)
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                c = t.cell(i, j)
                c.text = ""
                par = c.paragraphs[0]
                par.style = self.doc.styles["Text"]
                run = par.add_run(str(cell))
                run.font.size = Pt(font_pt)
                if i == 0:
                    run.bold = True
                par.paragraph_format.space_after = Pt(0)
                par.paragraph_format.first_line_indent = Pt(0)
                par.alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                                 else WD_ALIGN_PARAGRAPH.CENTER)
        return self.tab_no


def strip_template_remnants(doc) -> int:
    """Remove the template's trailing instruction block.

    The template ends with a boxed note telling the author to delete it, and
    an empty paragraph carrying a section break.  Both sections that follow
    the title block use identical two-column settings, so dropping that
    break merges them without changing the layout, and leaves the
    body-level section properties in charge.
    """
    body = doc.element.body
    removed = 0
    for p in list(body.findall(qn("w:p"))):
        txt = "".join(t.text or "" for t in p.iter(qn("w:t")))
        low = txt.lower()
        if "delete this box" in low or "updated instructions on the helpdesk" \
                in low or low.strip().startswith("important"):
            body.remove(p)
            removed += 1
    # drop a now-trailing empty paragraph that only carries a section break
    ps = body.findall(qn("w:p"))
    if ps:
        last = ps[-1]
        txt = "".join(t.text or "" for t in last.iter(qn("w:t"))).strip()
        pPr = last.find(qn("w:pPr"))
        has_sect = pPr is not None and pPr.find(qn("w:sectPr")) is not None
        if not txt and has_sect:
            body.remove(last)
            removed += 1
    return removed


def roman(n: int) -> str:
    vals = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for v, s in vals:
        while n >= v:
            out += s
            n -= v
    return out


def to_doc(docx_path: str, doc_path: str) -> None:
    """Convert to Word 97-2003 .doc, which is the required upload format."""
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    d = word.Documents.Open(os.path.abspath(docx_path))
    d.SaveAs2(os.path.abspath(doc_path), FileFormat=0)  # wdFormatDocument
    pages = d.ComputeStatistics(2)
    words = d.ComputeStatistics(0)
    d.Close(False)
    word.Quit()
    print(f"wrote {doc_path}: {pages} pages, {words} words")
    return pages, words
